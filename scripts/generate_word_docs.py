#!/usr/bin/env python3
"""Generate Word documents for participant onboarding forms."""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

DOWNLOADS_DIR = os.path.join(os.path.dirname(__file__), '..', 'downloads')

def add_heading(doc, text, level=0):
    """Add a heading to the document."""
    doc.add_heading(text, level=level)

def add_table(doc, headers, rows):
    """Add a table with field/value pairs."""
    table = doc.add_table(rows=len(rows)+1, cols=2)
    table.style = 'Table Grid'

    # Header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = headers[0]
    hdr_cells[1].text = headers[1]
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].bold = True

    # Data rows
    for i, row in enumerate(rows):
        row_cells = table.rows[i+1].cells
        row_cells[0].text = row[0]
        row_cells[1].text = row[1] if len(row) > 1 else ''

    doc.add_paragraph()  # Add spacing after table

def add_checkbox_list(doc, items):
    """Add a checkbox list."""
    for item in items:
        p = doc.add_paragraph()
        p.add_run('[ ] ').bold = True
        p.add_run(item)

def add_yes_no_table(doc, questions):
    """Add a table with Yes/No columns."""
    table = doc.add_table(rows=len(questions)+1, cols=3)
    table.style = 'Table Grid'

    # Header
    table.rows[0].cells[0].text = 'Question'
    table.rows[0].cells[1].text = 'Yes'
    table.rows[0].cells[2].text = 'No'
    for cell in table.rows[0].cells:
        cell.paragraphs[0].runs[0].bold = True

    for i, q in enumerate(questions):
        table.rows[i+1].cells[0].text = q
        table.rows[i+1].cells[1].text = '[ ]'
        table.rows[i+1].cells[2].text = '[ ]'

    doc.add_paragraph()

def create_onboarding_template():
    """Create the ISV Onboarding Document Template."""
    doc = Document()

    add_heading(doc, 'Onboarding Document Template', 0)
    doc.add_paragraph('Fill out this form and include it in your onboarding folder.')

    doc.add_paragraph().add_run('Required Documents:').bold = True
    doc.add_paragraph('Your onboarding folder must also include:')
    doc.add_paragraph('- Participant Agreement ([companyname] Participant Agreement)', style='List Bullet')
    doc.add_paragraph('- Contact Form ([companyname] Contact Form)', style='List Bullet')
    doc.add_paragraph('- Corporate Application ([companyname] Corporate Application) - if applying as a corporate entity', style='List Bullet')

    # Company Information
    add_heading(doc, 'Company Information', 1)
    add_table(doc, ['Field', 'Value'], [
        ['Company Legal Name', ''],
        ['Company Website', ''],
        ['Business Address', ''],
        ['Entity Type', '[ ] Individual  [ ] Corporate'],
    ])

    # Primary Technical Contact
    add_heading(doc, 'Primary Technical Contact', 1)
    add_table(doc, ['Field', 'Value'], [
        ['Full Name', ''],
        ['Title', ''],
        ['Email', ''],
        ['Phone', ''],
    ])

    # Secondary Technical Contact
    add_heading(doc, 'Secondary Technical Contact', 1)
    add_table(doc, ['Field', 'Value'], [
        ['Full Name', ''],
        ['Title', ''],
        ['Email', ''],
        ['Phone', ''],
    ])

    # API Access Details
    add_heading(doc, 'API Access Details', 1)
    add_table(doc, ['Field', 'Value'], [
        ['Requested Environment(s)', '[ ] Pre-production  [ ] Production'],
        ['Expected API Usage', '(e.g., orders/day, requests/minute)'],
        ['Static IP Addresses', '(list all IPs that will access the API)'],
    ])

    # Public Key Information
    add_heading(doc, 'Public Key Information', 1)

    doc.add_heading('Pre-Production Key (for testing)', 2)
    add_table(doc, ['Field', 'Value'], [
        ['Public Key Filename', '[companyname]_preprod_public_key.pem'],
        ['Key Generated Date', ''],
        ['Key Fingerprint', '(output of openssl command)'],
    ])

    doc.add_heading('Production Key', 2)
    add_table(doc, ['Field', 'Value'], [
        ['Public Key Filename', '[companyname]_prod_public_key.pem'],
        ['Key Generated Date', ''],
        ['Key Fingerprint', '(output of openssl command)'],
    ])

    # Use Case Description
    add_heading(doc, 'Use Case Description', 1)
    doc.add_paragraph('Describe how you plan to use the Polymarket Exchange API:')
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    # Acknowledgements
    add_heading(doc, 'Acknowledgements', 1)
    doc.add_paragraph('By submitting this onboarding request, we acknowledge that:')
    add_checkbox_list(doc, [
        'We have generated RSA key pairs for both preprod and production and will keep the private keys secure',
        'We will never share the private keys with anyone, including Polymarket',
        'We understand that if a private key is compromised, we must contact Polymarket immediately to rotate credentials',
        'We have reviewed the API documentation and understand the authentication flow',
        'We have completed and included the Participant Agreement',
        'We have completed and included the Contact Form',
    ])

    # Signature
    doc.add_paragraph()
    doc.add_paragraph('Submitted By: _________________________    Date: _____________')

    doc.save(os.path.join(DOWNLOADS_DIR, 'onboarding-template.docx'))
    print('Created: onboarding-template.docx')

def create_corporate_application():
    """Create the Corporate Application form."""
    doc = Document()

    add_heading(doc, 'Polymarket Exchange Corporate Application', 0)

    doc.add_paragraph().add_run('This form is required only for corporate entities.').bold = True
    doc.add_paragraph('Individual traders applying under their own name should skip this form.')
    doc.add_paragraph()
    doc.add_paragraph('Application Date: _________________')

    # 1. Corporate Entity Information
    add_heading(doc, '1. Corporate Entity Information', 1)
    add_table(doc, ['Field', 'Value'], [
        ['Legal Entity Name', ''],
        ['Entity Type', '[ ] C-Corp  [ ] S-Corp  [ ] LLC  [ ] Partnership  [ ] LP  [ ] LLP  [ ] Other: _______'],
        ['State/Country of Incorporation', ''],
        ['Date of Incorporation', ''],
        ['EIN / Tax ID Number', ''],
    ])

    # 2. Registered Agent
    add_heading(doc, '2. Registered Agent', 1)
    add_table(doc, ['Field', 'Value'], [
        ['Registered Agent Name', ''],
        ['Street Address', ''],
        ['City, State, Zip', ''],
        ['Phone Number', ''],
    ])

    # 3. Principal Place of Business
    add_heading(doc, '3. Principal Place of Business', 1)
    add_table(doc, ['Field', 'Value'], [
        ['Street Address', ''],
        ['City', ''],
        ['State/Province', ''],
        ['Postal Code', ''],
        ['Country', ''],
        ['Phone Number', ''],
    ])

    # 4. Ownership Structure
    add_heading(doc, '4. Ownership Structure', 1)

    doc.add_heading('4.1 Beneficial Owners', 2)
    doc.add_paragraph('List all individuals who own 25% or more of the entity, either directly or indirectly.')

    doc.add_paragraph().add_run('Owner 1:').bold = True
    add_table(doc, ['Field', 'Value'], [
        ['Full Legal Name', ''],
        ['Date of Birth', ''],
        ['SSN/Tax ID (last 4 digits)', ''],
        ['Ownership Percentage', '%'],
        ['Address', ''],
    ])

    doc.add_paragraph().add_run('Owner 2 (if applicable):').bold = True
    add_table(doc, ['Field', 'Value'], [
        ['Full Legal Name', ''],
        ['Date of Birth', ''],
        ['SSN/Tax ID (last 4 digits)', ''],
        ['Ownership Percentage', '%'],
        ['Address', ''],
    ])

    doc.add_paragraph().add_run('Owner 3 (if applicable):').bold = True
    add_table(doc, ['Field', 'Value'], [
        ['Full Legal Name', ''],
        ['Date of Birth', ''],
        ['SSN/Tax ID (last 4 digits)', ''],
        ['Ownership Percentage', '%'],
        ['Address', ''],
    ])

    doc.add_paragraph('Note: If no individual owns 25% or more, list the individual(s) with significant management responsibility (e.g., CEO, CFO, COO).')

    doc.add_heading('4.2 Control Person', 2)
    doc.add_paragraph('Individual with significant responsibility for managing the entity:')
    add_table(doc, ['Field', 'Value'], [
        ['Full Legal Name', ''],
        ['Title', ''],
        ['Date of Birth', ''],
        ['Address', ''],
    ])

    # 5. Officers and Directors
    add_heading(doc, '5. Officers and Directors', 1)

    doc.add_heading('Chief Executive Officer (CEO)', 2)
    add_table(doc, ['Field', 'Value'], [
        ['Full Legal Name', ''],
        ['Email Address', ''],
        ['Phone Number', ''],
    ])

    doc.add_heading('Chief Financial Officer (CFO)', 2)
    add_table(doc, ['Field', 'Value'], [
        ['Full Legal Name', ''],
        ['Email Address', ''],
        ['Phone Number', ''],
    ])

    doc.add_heading('Board of Directors / Managing Members', 2)
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = 'Name'
    table.rows[0].cells[1].text = 'Title'
    table.rows[0].cells[2].text = 'Email'
    for cell in table.rows[0].cells:
        cell.paragraphs[0].runs[0].bold = True
    doc.add_paragraph()

    # 6. Business Information
    add_heading(doc, '6. Business Information', 1)
    add_table(doc, ['Field', 'Value'], [
        ['Primary Business Activity', ''],
        ['Industry/Sector', ''],
        ['Years in Operation', ''],
        ['Number of Employees', ''],
        ['Annual Revenue Range', 'Under $1M / $1M-$10M / $10M-$50M / $50M-$100M / Over $100M'],
    ])

    doc.add_heading('Business Description', 2)
    doc.add_paragraph('Provide a brief description of your company\'s business activities:')
    doc.add_paragraph()
    doc.add_paragraph()

    # 7. Regulatory Status
    add_heading(doc, '7. Regulatory Status', 1)

    doc.add_heading('7.1 Licenses and Registrations', 2)
    doc.add_paragraph('Does the entity hold any financial services licenses or registrations?')
    doc.add_paragraph('[ ] Yes (complete table below)    [ ] No')

    table = doc.add_table(rows=3, cols=4)
    table.style = 'Table Grid'
    headers = ['License Type', 'Issuing Authority', 'License Number', 'Expiration Date']
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    doc.add_paragraph()

    doc.add_heading('7.2 Regulatory History', 2)
    doc.add_paragraph('Has the entity or any of its officers/directors ever been:')
    add_yes_no_table(doc, [
        'Subject to regulatory investigation or enforcement action?',
        'Denied a license or registration?',
        'Subject to a cease and desist order?',
        'Party to bankruptcy proceedings?',
    ])

    doc.add_paragraph('If "Yes" to any of the above, provide details:')
    doc.add_paragraph()
    doc.add_paragraph()

    # 8. Financial Information
    add_heading(doc, '8. Financial Information', 1)
    add_table(doc, ['Field', 'Value'], [
        ['Bank Name', ''],
        ['Bank Address', ''],
        ['Account Type', '[ ] Checking  [ ] Savings'],
        ['Account Number (last 4 digits)', ''],
        ['Routing Number', ''],
    ])

    doc.add_heading('Anticipated Trading Activity', 2)
    add_table(doc, ['Field', 'Value'], [
        ['Expected Monthly Trading Volume', '$'],
        ['Expected Number of End Users (for Partners)', ''],
        ['Primary Trading Strategy/Use Case', ''],
    ])

    # 9. Documents Required
    add_heading(doc, '9. Documents Required', 1)
    doc.add_paragraph('Please include the following documents with your application:')
    add_checkbox_list(doc, [
        'Certificate of Incorporation / Formation',
        'Articles of Organization / Operating Agreement',
        'Certificate of Good Standing (dated within 90 days)',
        'EIN Verification Letter (IRS CP-575 or equivalent)',
        'Government-issued ID for each beneficial owner',
        'Board Resolution authorizing API access (if applicable)',
    ])

    # 10. Certifications
    add_heading(doc, '10. Certifications', 1)
    doc.add_paragraph('By signing below, the undersigned certifies that:')
    add_checkbox_list(doc, [
        'All information provided in this application is true, accurate, and complete',
        'The entity is duly organized, validly existing, and in good standing',
        'The undersigned has authority to submit this application on behalf of the entity',
        'The entity will promptly notify Polymarket of any material changes to this information',
        'The entity complies with all applicable laws and regulations',
        'The entity has implemented adequate AML/KYC procedures (for Partners onboarding end users)',
    ])

    # 11. Signature
    add_heading(doc, '11. Signature', 1)
    doc.add_heading('Authorized Signatory', 2)
    add_table(doc, ['Field', 'Value'], [
        ['Signature', '_________________________'],
        ['Printed Name', ''],
        ['Title', ''],
        ['Date', ''],
    ])

    doc.add_paragraph()
    doc.add_paragraph('Corporate Seal / Stamp (if applicable):')
    doc.add_paragraph()

    doc.add_paragraph().add_run('Processing Time: ').bold = True
    doc.add_paragraph('Corporate applications typically require 5-10 business days for review. You will be contacted if additional documentation is needed.')

    doc.save(os.path.join(DOWNLOADS_DIR, 'corporate-application.docx'))
    print('Created: corporate-application.docx')

def create_contact_form():
    """Create the Contact Form."""
    doc = Document()

    add_heading(doc, 'Polymarket Exchange Contact Information Form', 0)

    doc.add_paragraph('This form collects contact information for your organization. Both contacts will be verified during the onboarding call and will receive credentials.')
    doc.add_paragraph()
    doc.add_paragraph('Submission Date: _________________')

    # 1. Organization Information
    add_heading(doc, '1. Organization Information', 1)
    add_table(doc, ['Field', 'Value'], [
        ['Legal Company Name', ''],
        ['DBA / Trade Name (if different)', ''],
        ['Company Website', ''],
        ['Company Phone Number', ''],
    ])

    doc.add_heading('Mailing Address', 2)
    add_table(doc, ['Field', 'Value'], [
        ['Street Address', ''],
        ['City', ''],
        ['State/Province', ''],
        ['Postal Code', ''],
        ['Country', ''],
    ])

    # 2. Primary Contact
    add_heading(doc, '2. Primary Contact', 1)
    doc.add_paragraph('This person will be the main point of contact for technical and operational matters.')
    add_table(doc, ['Field', 'Value'], [
        ['Full Legal Name', ''],
        ['Title/Position', ''],
        ['Email Address', ''],
        ['Phone Number', ''],
        ['Preferred Contact Method', '[ ] Email  [ ] Phone'],
        ['Time Zone', ''],
        ['Availability Hours', ''],
    ])

    # 3. Secondary Contact
    add_heading(doc, '3. Secondary Contact', 1)
    doc.add_paragraph('This person serves as backup and will also be verified during onboarding.')
    add_table(doc, ['Field', 'Value'], [
        ['Full Legal Name', ''],
        ['Title/Position', ''],
        ['Email Address', ''],
        ['Phone Number', ''],
        ['Preferred Contact Method', '[ ] Email  [ ] Phone'],
        ['Time Zone', ''],
        ['Availability Hours', ''],
    ])

    # 4. Billing Contact
    add_heading(doc, '4. Billing Contact', 1)
    doc.add_paragraph('(If different from Primary Contact)')
    add_table(doc, ['Field', 'Value'], [
        ['Full Legal Name', ''],
        ['Title/Position', ''],
        ['Email Address', ''],
        ['Phone Number', ''],
    ])

    doc.add_heading('Billing Address (if different from Mailing Address)', 2)
    add_table(doc, ['Field', 'Value'], [
        ['Street Address', ''],
        ['City', ''],
        ['State/Province', ''],
        ['Postal Code', ''],
        ['Country', ''],
    ])

    # 5. Technical Contact
    add_heading(doc, '5. Technical Contact', 1)
    doc.add_paragraph('(If different from Primary Contact)')
    add_table(doc, ['Field', 'Value'], [
        ['Full Legal Name', ''],
        ['Title/Position', ''],
        ['Email Address', ''],
        ['Phone Number', ''],
        ['GitHub Username (optional)', ''],
    ])

    # 6. Emergency Contact
    add_heading(doc, '6. Emergency Contact', 1)
    doc.add_paragraph('For urgent security or operational issues outside business hours.')
    add_table(doc, ['Field', 'Value'], [
        ['Full Legal Name', ''],
        ['Title/Position', ''],
        ['Email Address', ''],
        ['Phone Number', ''],
        ['Preferred Contact Method', '[ ] Email  [ ] Phone  [ ] SMS'],
    ])

    # 7. Notification Preferences
    add_heading(doc, '7. Notification Preferences', 1)
    doc.add_paragraph('How should we contact you for different types of communications?')

    table = doc.add_table(rows=6, cols=4)
    table.style = 'Table Grid'
    headers = ['Communication Type', 'Email', 'Phone', 'SMS']
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    comm_types = ['API Status Updates', 'Security Alerts', 'Maintenance Notifications', 'Product Updates', 'Billing/Invoices']
    for i, ct in enumerate(comm_types):
        table.rows[i+1].cells[0].text = ct
        table.rows[i+1].cells[1].text = '[ ]'
        table.rows[i+1].cells[2].text = '[ ]'
        table.rows[i+1].cells[3].text = '[ ]'
    doc.add_paragraph()

    # 8. Distribution Lists
    add_heading(doc, '8. Distribution Lists', 1)
    doc.add_paragraph('Provide any shared email addresses for team communications:')

    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = 'Purpose'
    table.rows[0].cells[1].text = 'Email Address'
    for cell in table.rows[0].cells:
        cell.paragraphs[0].runs[0].bold = True

    purposes = ['Technical/Engineering', 'Operations', 'Compliance', 'Executive']
    for i, p in enumerate(purposes):
        table.rows[i+1].cells[0].text = p
    doc.add_paragraph()

    # 9. Acknowledgements
    add_heading(doc, '9. Acknowledgements', 1)
    add_checkbox_list(doc, [
        'I confirm all contact information provided is accurate',
        'I authorize Polymarket to contact the individuals listed for onboarding and operational purposes',
        'I will update this form if any contact information changes',
        'Both Primary and Secondary contacts will be present for the verification call',
    ])

    # Signature
    doc.add_paragraph()
    doc.add_paragraph('Submitted By: _________________________    Date: _____________')
    doc.add_paragraph('Title: _________________________')

    doc.add_paragraph()
    doc.add_paragraph().add_run('Verification Call Reminder: ').bold = True
    doc.add_paragraph('Both Primary and Secondary contacts must be present during the onboarding verification call with government-issued photo ID.')

    doc.save(os.path.join(DOWNLOADS_DIR, 'contact-form.docx'))
    print('Created: contact-form.docx')

def create_participant_agreement():
    """Create the Participant Agreement."""
    doc = Document()

    add_heading(doc, 'Polymarket Exchange API Participant Agreement', 0)

    doc.add_paragraph('This agreement establishes the terms under which you will access and use the Polymarket Exchange API.')
    doc.add_paragraph()
    doc.add_paragraph('Effective Date: _________________')
    doc.add_paragraph()
    doc.add_paragraph().add_run('Between:').bold = True
    doc.add_paragraph('Polymarket US, LLC ("Polymarket")')
    doc.add_paragraph('and')
    doc.add_paragraph('Participant (as identified below)')

    # 1. Participant Information
    add_heading(doc, '1. Participant Information', 1)
    add_table(doc, ['Field', 'Value'], [
        ['Legal Name', ''],
        ['Entity Type', '[ ] Individual  [ ] Corporation  [ ] LLC  [ ] Partnership  [ ] Other: _______'],
        ['Jurisdiction of Formation', ''],
        ['Principal Business Address', ''],
        ['EIN/Tax ID', ''],
    ])

    # 2. Authorized Representatives
    add_heading(doc, '2. Authorized Representatives', 1)

    doc.add_heading('Primary Authorized Representative', 2)
    add_table(doc, ['Field', 'Value'], [
        ['Full Legal Name', ''],
        ['Title/Position', ''],
        ['Email Address', ''],
        ['Phone Number', ''],
    ])

    doc.add_heading('Secondary Authorized Representative', 2)
    add_table(doc, ['Field', 'Value'], [
        ['Full Legal Name', ''],
        ['Title/Position', ''],
        ['Email Address', ''],
        ['Phone Number', ''],
    ])

    # 3. Participant Type
    add_heading(doc, '3. Participant Type', 1)
    doc.add_paragraph('Select all that apply:')
    add_checkbox_list(doc, [
        'Direct Trader - Trading on own behalf using own capital',
        'Retail Partner (ISV) - Building a platform for retail end-users',
        'Introducing Broker (IB) - Introducing clients to Polymarket',
        'Futures Commission Merchant (FCM) - Licensed FCM',
    ])

    # 4. Representations and Warranties
    add_heading(doc, '4. Representations and Warranties', 1)
    doc.add_paragraph('By signing this Agreement, Participant represents and warrants that:')

    doc.add_heading('4.1 Legal Authority', 2)
    add_checkbox_list(doc, [
        'Participant has full legal authority to enter into this Agreement',
        'The individual signing has authority to bind Participant to this Agreement',
        'Participant is not subject to any legal or regulatory restriction that would prohibit participation',
    ])

    doc.add_heading('4.2 Regulatory Compliance', 2)
    add_checkbox_list(doc, [
        'Participant will comply with all applicable laws and regulations',
        'Participant maintains all required licenses for its business activities',
        'Participant will immediately notify Polymarket of any regulatory inquiry or action',
    ])

    doc.add_heading('4.3 Financial Standing', 2)
    add_checkbox_list(doc, [
        'Participant is not insolvent, bankrupt, or subject to insolvency proceedings',
        'Participant has adequate capital to meet its anticipated trading obligations',
    ])

    doc.add_heading('4.4 Technical Capabilities', 2)
    add_checkbox_list(doc, [
        'Participant has technical capability to securely integrate with the API',
        'Participant will maintain security of all credentials and private keys',
        'Participant will implement appropriate access controls and monitoring',
    ])

    # 5. Obligations
    add_heading(doc, '5. Obligations', 1)

    doc.add_heading('5.1 Security Obligations', 2)
    add_checkbox_list(doc, [
        'Private keys will never be shared with any third party, including Polymarket',
        'Participant will immediately report any security breach or key compromise',
        'Participant will implement industry-standard security practices',
    ])

    doc.add_heading('5.2 Operational Obligations', 2)
    add_checkbox_list(doc, [
        'Participant will comply with API rate limits and usage policies',
        'Participant will maintain accurate records of all API activity',
        'Participant will cooperate with Polymarket in investigating any issues',
    ])

    doc.add_heading('5.3 Reporting Obligations', 2)
    add_checkbox_list(doc, [
        'Participant will promptly report any material changes to information provided',
        'Participant will notify Polymarket of changes to authorized representatives',
        'Participant will provide additional information reasonably requested by Polymarket',
    ])

    # 6. Acknowledgements
    add_heading(doc, '6. Acknowledgements', 1)
    doc.add_paragraph('By signing below, Participant acknowledges and agrees:')
    add_checkbox_list(doc, [
        'I have read and understand the Polymarket Exchange API documentation',
        'I understand the risks associated with trading on prediction markets',
        'I understand that Polymarket may suspend or terminate API access at any time',
        'I consent to Polymarket\'s data collection and privacy practices',
        'I will comply with the Polymarket Terms of Service and API Usage Policy',
    ])

    # 7. Signatures
    add_heading(doc, '7. Signatures', 1)

    doc.add_heading('For Participant:', 2)
    add_table(doc, ['Field', 'Value'], [
        ['Signature', '_________________________'],
        ['Printed Name', ''],
        ['Title', ''],
        ['Date', ''],
    ])

    doc.add_heading('For Polymarket (Office Use Only):', 2)
    add_table(doc, ['Field', 'Value'], [
        ['Signature', '_________________________'],
        ['Printed Name', ''],
        ['Title', ''],
        ['Date', ''],
    ])

    doc.add_paragraph()
    doc.add_paragraph().add_run('Important: ').bold = True
    doc.add_paragraph('This document must be completed in full, signed by an authorized representative, and included in your onboarding folder. Incomplete agreements will delay the onboarding process.')

    doc.save(os.path.join(DOWNLOADS_DIR, 'participant-agreement.docx'))
    print('Created: participant-agreement.docx')

def create_getting_started_onboarding():
    """Create the Getting Started Onboarding Template (simpler version)."""
    doc = Document()

    add_heading(doc, 'Partner Onboarding Document Template', 0)
    doc.add_paragraph('Fill out this form and include it in your shared onboarding folder.')

    # Company Information
    add_heading(doc, 'Company Information', 1)
    add_table(doc, ['Field', 'Value'], [
        ['Company Legal Name', ''],
        ['Company Website', ''],
        ['Business Address', ''],
    ])

    # Primary Technical Contact
    add_heading(doc, 'Primary Technical Contact', 1)
    add_table(doc, ['Field', 'Value'], [
        ['Full Name', ''],
        ['Title', ''],
        ['Email', ''],
        ['Phone', ''],
    ])

    # Secondary Technical Contact
    add_heading(doc, 'Secondary Technical Contact', 1)
    add_table(doc, ['Field', 'Value'], [
        ['Full Name', ''],
        ['Title', ''],
        ['Email', ''],
        ['Phone', ''],
    ])

    # API Access Details
    add_heading(doc, 'API Access Details', 1)
    add_table(doc, ['Field', 'Value'], [
        ['Requested Environment(s)', '[ ] Development  [ ] Pre-production  [ ] Production'],
        ['Expected API Usage', '(e.g., orders/day, requests/minute)'],
        ['Static IP Addresses', '(list all IPs that will access the API)'],
    ])

    # Public Key Information
    add_heading(doc, 'Public Key Information', 1)

    doc.add_heading('Pre-Production Key (for testing)', 2)
    add_table(doc, ['Field', 'Value'], [
        ['Public Key Filename', '[firmname]_preprod_public_key.pem'],
        ['Key Generated Date', ''],
        ['Key Fingerprint', '(output of openssl command)'],
    ])

    doc.add_heading('Production Key', 2)
    add_table(doc, ['Field', 'Value'], [
        ['Public Key Filename', '[firmname]_prod_public_key.pem'],
        ['Key Generated Date', ''],
        ['Key Fingerprint', '(output of openssl command)'],
    ])

    # Use Case Description
    add_heading(doc, 'Use Case Description', 1)
    doc.add_paragraph('Describe how you plan to use the Polymarket Exchange API:')
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    # Acknowledgements
    add_heading(doc, 'Acknowledgements', 1)
    doc.add_paragraph('By submitting this onboarding request, we acknowledge that:')
    add_checkbox_list(doc, [
        'We have generated an RSA key pair and will keep the private key secure',
        'We will never share the private key with anyone, including Polymarket',
        'We understand that if the private key is compromised, we must contact Polymarket immediately to rotate credentials',
        'We have reviewed the API documentation and understand the authentication flow',
    ])

    # Signature
    doc.add_paragraph()
    doc.add_paragraph('Submitted By: _________________________    Date: _____________')

    doc.save(os.path.join(DOWNLOADS_DIR, 'partner-onboarding-template.docx'))
    print('Created: partner-onboarding-template.docx')

if __name__ == '__main__':
    os.makedirs(DOWNLOADS_DIR, exist_ok=True)
    create_onboarding_template()
    create_corporate_application()
    create_contact_form()
    create_participant_agreement()
    create_getting_started_onboarding()
    print('\nAll Word documents created successfully!')
